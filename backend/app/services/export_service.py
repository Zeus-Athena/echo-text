"""
Export Service
导出服务 - 支持 PDF、Markdown、Word、SRT 格式导出
"""

from dataclasses import dataclass
from io import BytesIO

from app.models.recording import Recording


@dataclass
class ExportOptions:
    """导出选项"""

    include_transcript: bool = True
    include_translation: bool = True
    include_summary: bool = True
    include_timestamps: bool = True


class ExportService:
    """文档导出服务"""

    def __init__(self, recording: Recording):
        self.recording = recording

    async def export_markdown(self, options: ExportOptions = None) -> str:
        """
        导出为 Markdown 格式
        """
        options = options or ExportOptions()
        lines = []

        # Title
        lines.append(f"# {self.recording.title}")
        lines.append("")

        # Metadata
        lines.append("## 录音信息")
        lines.append("")
        lines.append(f"- **录音时长**: {self._format_duration(self.recording.duration_seconds)}")
        lines.append(f"- **源语言**: {self.recording.source_lang}")
        lines.append(f"- **目标语言**: {self.recording.target_lang}")
        lines.append(f"- **创建时间**: {self.recording.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append("")

        # AI Summary
        if options.include_summary and self.recording.ai_summary:
            summary = self.recording.ai_summary
            lines.append("## AI 总结")
            lines.append("")
            if summary.summary:
                lines.append(summary.summary)
                lines.append("")

            if summary.key_points:
                lines.append("### 要点")
                lines.append("")
                for point in summary.key_points:
                    lines.append(f"- {point}")
                lines.append("")

            if summary.action_items:
                lines.append("### 待办事项")
                lines.append("")
                for item in summary.action_items:
                    lines.append(f"- [ ] {item}")
                lines.append("")

        # Transcript
        if options.include_transcript and self.recording.transcript:
            lines.append("## 转录内容")
            lines.append("")
            if options.include_timestamps and self.recording.transcript.segments:
                for seg in self.recording.transcript.segments:
                    timestamp = self._format_timestamp(seg.get("start", 0))
                    text = seg.get("text", "")
                    speaker = seg.get("speaker", "")
                    if speaker:
                        lines.append(f"**[{timestamp}] {speaker}:** {text}")
                    else:
                        lines.append(f"**[{timestamp}]** {text}")
                lines.append("")
            else:
                lines.append(self.recording.transcript.full_text or "")
                lines.append("")

        # Translation
        if options.include_translation and self.recording.translation:
            lines.append("## 翻译内容")
            lines.append("")
            if options.include_timestamps and self.recording.translation.segments:
                for seg in self.recording.translation.segments:
                    timestamp = self._format_timestamp(seg.get("start", 0))
                    text = seg.get("text", "")
                    lines.append(f"**[{timestamp}]** {text}")
                lines.append("")
            else:
                lines.append(self.recording.translation.full_text or "")
                lines.append("")

        return "\n".join(lines)

    async def export_srt(self, use_translation: bool = False) -> str:
        """
        导出为 SRT 字幕格式

        Args:
            use_translation: True 使用翻译内容，False 使用转录内容
        """
        segments = []

        if use_translation and self.recording.translation and self.recording.translation.segments:
            segments = self.recording.translation.segments
        elif self.recording.transcript and self.recording.transcript.segments:
            segments = self.recording.transcript.segments

        if not segments:
            # 如果没有分段，创建单个字幕
            text = ""
            if use_translation and self.recording.translation:
                text = self.recording.translation.full_text or ""
            elif self.recording.transcript:
                text = self.recording.transcript.full_text or ""

            if text:
                return f"1\n00:00:00,000 --> {self._format_srt_time(self.recording.duration_seconds)}\n{text}\n"
            return ""

        lines = []
        for i, seg in enumerate(segments, 1):
            start = seg.get("start", 0)
            end = seg.get("end", start + 1)
            text = seg.get("text", "").strip()
            speaker = seg.get("speaker", "")

            if not text:
                continue

            if speaker:
                text = f"[{speaker}] {text}"

            lines.append(str(i))
            lines.append(f"{self._format_srt_time(start)} --> {self._format_srt_time(end)}")
            lines.append(text)
            lines.append("")

        return "\n".join(lines)

    async def export_pdf(self, options: ExportOptions = None) -> bytes:
        """
        导出为 PDF 格式
        使用 reportlab 生成 PDF
        """
        try:
            from reportlab.lib.enums import TA_LEFT
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import cm
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.platypus import (
                Paragraph,
                SimpleDocTemplate,
                Spacer,
            )
        except ImportError:
            raise ImportError("PDF 导出需要安装 reportlab 库: pip install reportlab") from None

        options = options or ExportOptions()
        buffer = BytesIO()

        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
        )

        # 注册中文字体 (使用系统字体或回退)
        chinese_font = "Helvetica"  # 默认回退
        font_paths = [
            # macOS 字体
            ("/System/Library/Fonts/STHeiti Light.ttc", 0),
            ("/System/Library/Fonts/STHeiti Medium.ttc", 0),
            ("/Library/Fonts/Arial Unicode.ttf", None),
            ("/System/Library/Fonts/Supplemental/Arial Unicode.ttf", None),
            # Linux 字体
            ("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc", 0),
            ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", 0),
            ("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc", 0),
            # Windows 字体
            ("C:/Windows/Fonts/msyh.ttc", 0),  # 微软雅黑
            ("C:/Windows/Fonts/simsun.ttc", 0),  # 宋体
        ]

        import os

        for font_path, subfont_index in font_paths:
            if os.path.exists(font_path):
                try:
                    if subfont_index is not None and font_path.endswith(".ttc"):
                        pdfmetrics.registerFont(
                            TTFont("Chinese", font_path, subfontIndex=subfont_index)
                        )
                    else:
                        pdfmetrics.registerFont(TTFont("Chinese", font_path))
                    chinese_font = "Chinese"
                    break
                except Exception:
                    continue

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "Title", parent=styles["Title"], fontName=chinese_font, fontSize=18, spaceAfter=20
        )
        heading_style = ParagraphStyle(
            "Heading", parent=styles["Heading2"], fontName=chinese_font, fontSize=14, spaceAfter=10
        )
        body_style = ParagraphStyle(
            "Body",
            parent=styles["Normal"],
            fontName=chinese_font,
            fontSize=11,
            leading=16,
            alignment=TA_LEFT,
        )

        story = []

        # Title
        story.append(Paragraph(self.recording.title, title_style))
        story.append(Spacer(1, 10))

        # Metadata
        meta_text = f"""
        <b>录音时长:</b> {self._format_duration(self.recording.duration_seconds)}<br/>
        <b>源语言:</b> {self.recording.source_lang}<br/>
        <b>目标语言:</b> {self.recording.target_lang}<br/>
        <b>创建时间:</b> {self.recording.created_at.strftime("%Y-%m-%d %H:%M:%S")}
        """
        story.append(Paragraph(meta_text, body_style))
        story.append(Spacer(1, 20))

        # AI Summary
        if options.include_summary and self.recording.ai_summary:
            story.append(Paragraph("AI 总结", heading_style))
            if self.recording.ai_summary.summary:
                story.append(Paragraph(self.recording.ai_summary.summary, body_style))
            story.append(Spacer(1, 10))

        # Transcript
        if options.include_transcript and self.recording.transcript:
            story.append(Paragraph("转录内容", heading_style))
            text = self.recording.transcript.full_text or ""
            # 分段处理长文本
            for para in text.split("\n"):
                if para.strip():
                    story.append(Paragraph(para, body_style))
            story.append(Spacer(1, 10))

        # Translation
        if options.include_translation and self.recording.translation:
            story.append(Paragraph("翻译内容", heading_style))
            text = self.recording.translation.full_text or ""
            for para in text.split("\n"):
                if para.strip():
                    story.append(Paragraph(para, body_style))

        doc.build(story)
        buffer.seek(0)
        return buffer.read()

    async def export_docx(self, options: ExportOptions = None) -> bytes:
        """
        导出为 Word (docx) 格式
        使用 python-docx 生成文档
        """
        try:
            from docx import Document
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        except ImportError:
            raise ImportError("Word 导出需要安装 python-docx 库: pip install python-docx") from None

        options = options or ExportOptions()
        doc = Document()

        # Title
        title = doc.add_heading(self.recording.title, 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Metadata
        doc.add_heading("录音信息", level=1)
        meta = doc.add_paragraph()
        meta.add_run("录音时长: ").bold = True
        meta.add_run(f"{self._format_duration(self.recording.duration_seconds)}\n")
        meta.add_run("源语言: ").bold = True
        meta.add_run(f"{self.recording.source_lang}\n")
        meta.add_run("目标语言: ").bold = True
        meta.add_run(f"{self.recording.target_lang}\n")
        meta.add_run("创建时间: ").bold = True
        meta.add_run(f"{self.recording.created_at.strftime('%Y-%m-%d %H:%M:%S')}")

        # AI Summary
        if options.include_summary and self.recording.ai_summary:
            doc.add_heading("AI 总结", level=1)
            if self.recording.ai_summary.summary:
                doc.add_paragraph(self.recording.ai_summary.summary)

            if self.recording.ai_summary.key_points:
                doc.add_heading("要点", level=2)
                for point in self.recording.ai_summary.key_points:
                    doc.add_paragraph(f"• {point}")

            if self.recording.ai_summary.action_items:
                doc.add_heading("待办事项", level=2)
                for item in self.recording.ai_summary.action_items:
                    doc.add_paragraph(f"☐ {item}")

        # Transcript
        if options.include_transcript and self.recording.transcript:
            doc.add_heading("转录内容", level=1)
            if options.include_timestamps and self.recording.transcript.segments:
                for seg in self.recording.transcript.segments:
                    timestamp = self._format_timestamp(seg.get("start", 0))
                    text = seg.get("text", "")
                    speaker = seg.get("speaker", "")
                    p = doc.add_paragraph()
                    p.add_run(f"[{timestamp}] ").bold = True
                    if speaker:
                        p.add_run(f"{speaker}: ").italic = True
                    p.add_run(text)
            else:
                doc.add_paragraph(self.recording.transcript.full_text or "")

        # Translation
        if options.include_translation and self.recording.translation:
            doc.add_heading("翻译内容", level=1)
            if options.include_timestamps and self.recording.translation.segments:
                for seg in self.recording.translation.segments:
                    timestamp = self._format_timestamp(seg.get("start", 0))
                    text = seg.get("text", "")
                    p = doc.add_paragraph()
                    p.add_run(f"[{timestamp}] ").bold = True
                    p.add_run(text)
            else:
                doc.add_paragraph(self.recording.translation.full_text or "")

        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def _format_duration(self, seconds: int) -> str:
        """格式化时长"""
        hours = seconds // 3600
        minutes = (seconds % 3600) // 60
        secs = seconds % 60
        if hours > 0:
            return f"{hours}:{minutes:02d}:{secs:02d}"
        return f"{minutes}:{secs:02d}"

    def _format_timestamp(self, seconds: float) -> str:
        """格式化时间戳 (MM:SS)"""
        minutes = int(seconds // 60)
        secs = int(seconds % 60)
        return f"{minutes:02d}:{secs:02d}"

    def _format_srt_time(self, seconds: float) -> str:
        """格式化 SRT 时间戳 (HH:MM:SS,mmm)"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"
